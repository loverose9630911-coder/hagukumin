# -*- coding: utf-8 -*-
"""build.py — 図版生成 → Word(.docx)組版 → PDF変換 を一括実行する。

出力先: manuals/
  00_共通編_Claude×MCP連携の基礎.docx / .pdf
  01〜10 各サービス編 .docx / .pdf
  README.md(索引)
"""
import os
import subprocess
import genfig
import scenes
from content import COMMON, OVERVIEW_TABLE, SERVICES, DISCLAIMER

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), "..")
OUT = os.path.join(ROOT, "manuals")
FIG = os.path.join(OUT, "figures")
os.makedirs(FIG, exist_ok=True)

ASCII_FONT = "Arial"
CJK_FONT = "IPAGothic"   # 本環境(LibreOffice)に実在。Wordでは日本語フォントへ自動代替される。
ACCENT = RGBColor(0x25, 0x5F, 0xB5)
HEADER = RGBColor(0x2D, 0x37, 0x48)
GREY = RGBColor(0x6C, 0x75, 0x7D)


# ---------------------------------------------------------------- docx helpers
def _set_cjk(run):
    run.font.name = ASCII_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def _shade(el, fill):
    pr = el.get_or_add_pPr() if el.tag.endswith("}p") else el
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _para_border(p, color="D97706", sz="18", space="8"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), space)
        e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def run(p, text, size=10.5, bold=False, color=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    _set_cjk(r)
    return r


def para(doc, text="", size=10.5, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, size, bold, color)
    return p


def h1(doc, text, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    label = f"{num}. {text}" if num else text
    run(p, label, 15.5, True, HEADER)
    # 下線(段落罫線 bottom)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "10")
    e.set(qn("w:space"), "4"); e.set(qn("w:color"), "255FB5")
    pbdr.append(e); pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run(p, "■ " + text, 12, True, ACCENT)
    return p


def bullets(doc, items, ordered=False):
    for i, it in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        mark = f"{i+1}. " if ordered else "・"
        run(p, mark + it, 10.5)


def figure(doc, path, caption, width_cm=16.4):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    run(cp, caption, 9, False, GREY)


def placeholder_block(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    _shade(p._p.get_or_add_pPr(), "FFFBEB")
    _para_border(p)
    run(p, "［ 実スクリーンショット差込欄 ］", 10.5, True, RGBColor(0xD9, 0x77, 0x06))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    _shade(p2._p.get_or_add_pPr(), "FFFBEB")
    _para_border(p2)
    run(p2, "▼ ここに実画面のスクリーンショットを貼り付け: " + caption, 9.5, False, GREY)


def callout_box(doc, title, lines, fill="EFF6FF", border="255FB5", tcol=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    _shade(p._p.get_or_add_pPr(), fill)
    _para_border(p, color=border, sz="8")
    run(p, title, 10.5, True, tcol)
    for ln in lines:
        pp = doc.add_paragraph()
        _shade(pp._p.get_or_add_pPr(), fill)
        _para_border(pp, color=border, sz="8")
        pp.paragraph_format.space_after = Pt(2)
        run(pp, ln, 10)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


def code_box(doc, lines):
    for i, ln in enumerate(lines):
        p = doc.add_paragraph()
        _shade(p._p.get_or_add_pPr(), "1E222D")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.size = Pt(9.5)
        r.font.name = "Consolas"
        r.font.color.rgb = RGBColor(0xDC, 0xDF, 0xE4)
        _set_cjk(r)
        r.font.name = "Consolas"
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


def setup_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = ASCII_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    for m in ("top_margin", "bottom_margin"):
        setattr(sec, m, Cm(2.0))
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Cm(2.2))
    return doc


def cover(doc, title, subtitle, tag=""):
    for _ in range(2):
        doc.add_paragraph()
    para(doc, "営業ツール × Claude 連携マニュアル", 13, True, ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, title, 22, True, HEADER, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    if subtitle:
        para(doc, subtitle, 12.5, False, GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    if tag:
        para(doc, tag, 11, True, ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    # 注意書きボックス
    callout_box(doc, "本マニュアルについて(必ずお読みください)", [DISCLAIMER],
                fill="FFF7ED", border="D97706", tcol=RGBColor(0xB4, 0x53, 0x09))
    para(doc, "作成日: 2026年6月28日 / 基準記事: HubSpot『国内の主要SFAツール10選』",
         9.5, False, GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def legend(doc):
    callout_box(doc, "図の見方", [
        "赤い枠 … 操作する対象(クリック/入力する箇所)を示します。",
        "番号バッジ ①②③ … 操作の順序と、右側の注釈に対応します。",
        "オレンジの破線枠『実スクリーンショット差込欄』… 実環境で撮影した画面を貼り付ける場所です。",
    ], fill="F8FAFC", border="CED4DA", tcol=HEADER)


# ---------------------------------------------------------------- PDF
def to_pdf(docx_path):
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", OUT, docx_path],
                   check=True, capture_output=True, timeout=180)


# ---------------------------------------------------------------- 共通編
def build_common():
    doc = setup_doc()
    cover(doc, COMMON["title"], COMMON["subtitle"])
    doc.add_page_break()

    h1(doc, "はじめに", 1)
    para(doc, COMMON["intro"])
    legend(doc)

    h1(doc, "MCP(Model Context Protocol)とは", 2)
    para(doc, COMMON["mcp_what"])
    fa = genfig.arch_diagram(os.path.join(FIG, "common_arch.png"),
                             "SFA / CRM", "official",
                             "公式MCPが無いサービスは iPaaS(Zapier) またはカスタムMCPで連携します。")
    figure(doc, fa, "図1: Claude × MCP × SFA/CRM の基本構成")

    h1(doc, "連携の3つのパターン", 3)
    for t, d in COMMON["patterns"]:
        h2(doc, t)
        para(doc, d)
    ff = genfig.flow_diagram(os.path.join(FIG, "common_flow.png"), "導入の基本ステップ",
                             [("方式を選ぶ", "公式MCP/iPaaS/カスタム"),
                              ("認証を準備", "トークン/OAuth/APIキー"),
                              ("Claudeに登録", "コネクタ or config.json"),
                              ("動作確認", "参照系でテスト"),
                              ("活用開始", "業務で運用")])
    figure(doc, ff, "図2: 導入の基本ステップ")

    h1(doc, "Claude側の準備", 4)
    for t, d in COMMON["client_prep"]:
        h2(doc, t)
        para(doc, d)
    code_box(doc, [
        "// Claude Desktop: claude_desktop_config.json の例",
        "{",
        '  "mcpServers": {',
        '    "<サービス名>": {',
        '      "command": "npx",',
        '      "args": ["-y", "<MCPサーバーのパッケージ>"],',
        '      "env": { "API_TOKEN": "********" }',
        "    }",
        "  }",
        "}",
    ])
    placeholder_block(doc, "実際の設定ファイル編集画面/コネクタ追加画面")

    h1(doc, "各サービスの連携方式 早見表", 5)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["サービス", "提供元", "連携方式", "公式MCP", "主な接続方法"]):
        c.paragraphs[0].clear()
        run(c.paragraphs[0], txt, 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
        _shade(c._tc.get_or_add_tcPr(), "2D3748")
    for row in OVERVIEW_TABLE:
        cells = tbl.add_row().cells
        for c, txt in zip(cells, row):
            c.paragraphs[0].clear()
            run(c.paragraphs[0], txt, 9)
    para(doc, "")
    para(doc, "※ 公式MCPの有無・仕様は変更される場合があります。最新情報は各社公式をご確認ください。",
         9, False, GREY)

    h1(doc, "セキュリティ・運用上の注意", 6)
    bullets(doc, COMMON["security"])
    callout_box(doc, "推奨: スモールスタート", [
        "①参照のみの権限で接続 → ②少人数で試験運用 → ③書き込み権限を段階的に付与 → ④全体展開、",
        "の順で進めると、事故を防ぎつつ定着させやすくなります。",
    ])

    path = os.path.join(OUT, "00_共通編_Claude×MCP連携の基礎.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------- サービス編
def build_service(s):
    doc = setup_doc()
    cover(doc, s["name"], s["tagline"], tag=f"提供元: {s['vendor']}")
    doc.add_page_break()
    legend(doc)

    h1(doc, "本サービスと Claude 連携の概要", 1)
    para(doc, s["summary"])

    h1(doc, "連携方式と構成", 2)
    para(doc, s["method"])
    arch = genfig.arch_diagram(os.path.join(FIG, f"{s['file']}_arch.png"),
                               s["name"], s["mode"])
    figure(doc, arch, f"図1: {s['name']} と Claude の連携構成図")
    flow = genfig.flow_diagram(os.path.join(FIG, f"{s['file']}_flow.png"),
                               "連携の流れ", s["steps"])
    figure(doc, flow, "図2: 導入ステップの全体像")

    h1(doc, "事前準備(前提条件)", 3)
    bullets(doc, s["prereq"])

    h1(doc, "連携手順", 4)
    # 手順1: 認証情報の発行(token scene)
    if s.get("token"):
        h2(doc, "手順1: サービス側で認証情報を発行する")
        t = s["token"]
        tk = scenes.scene_token(
            s["name"], f"{s['file']}_token.png",
            f"図3: {s['name']} で認証情報を発行する",
            t["panel_title"], t["rows"], t["annot_row"], t["callouts"], t["ph"])
        figure(doc, tk, f"図3: {t['panel_title']}(操作対象を赤枠で表示)")
        placeholder_block(doc, t["ph"])

    # 手順2: Claudeへ登録(connect scene)
    h2(doc, "手順2: Claude にMCPを登録する")
    ct = s["connect"]
    if ct == "connector":
        cp = s["connect_params"]
        cf = scenes.scene_connector(
            s["name"], f"{s['file']}_connect.png",
            f"図4: Claude のコネクタに {s['name']} を追加する",
            cp["url"], cp["steps"], cp["ph"])
        figure(doc, cf, "図4: Claudeのコネクタ追加画面(公式リモートMCP)")
        placeholder_block(doc, cp["ph"])
    elif ct == "config":
        cp = s["config_params"]
        cf = scenes.scene_config(
            s["name"], f"{s['file']}_connect.png",
            f"図4: claude_desktop_config.json に {s['name']} のMCPを登録する",
            cp["pkg"], cp["args"], cp["envs"], cp["ph"], cp.get("note", ""))
        figure(doc, cf, "図4: Claude Desktop 設定ファイルへの登録")
        placeholder_block(doc, cp["ph"])
    else:  # zapier
        cp = s["zapier_params"]
        cf = scenes.scene_zapier(
            s["name"], f"{s['file']}_connect.png",
            f"図4: Zapier(iPaaS)経由で {s['name']} を連携する", cp["ph"])
        figure(doc, cf, "図4: Zapier MCP の設定画面(iPaaS経由)")
        placeholder_block(doc, cp["ph"])

    h2(doc, "手順3: 動作確認(参照系でテスト)")
    para(doc, "まずは『データを5件表示して』のような参照のみの指示で接続を確認します。"
              "正しく取得できれば連携は成功です。続いて作成・更新の操作に進みます。")

    h1(doc, "活用例・使用例", 5)
    para(doc, "実際のチャットでの活用イメージを示します。自然言語で指示すると、ClaudeがMCPツールを"
              "選んで実行し、結果を要約して回答します。")
    # 代表ユースケースをチャット図に
    uc0 = s["use_cases"][0]
    chat = scenes.scene_chat(
        s["name"], f"{s['file']}_chat.png",
        f"図5: チャットでの活用例 ―『{uc0['title']}』",
        uc0["prompt"], uc0["tool"], uc0["reply"],
        "実際のClaudeチャット画面(指示と回答)を貼り付け")
    figure(doc, chat, f"図5: 活用例 ―『{uc0['title']}』のチャット例")
    placeholder_block(doc, "実際のClaudeチャット画面(指示と回答)を貼り付け")

    for i, uc in enumerate(s["use_cases"]):
        h2(doc, f"活用例{i+1}: {uc['title']}")
        callout_box(doc, "プロンプト例(Claudeへの指示)", ['「' + uc["prompt"] + '」'],
                    fill="EFF6FF", border="255FB5")
        para(doc, "▶ 期待される動作: " + uc["desc"], 10)
        para(doc, "▶ 呼び出されるMCPツール(例): " + uc["tool"], 9.5, color=GREY)

    h1(doc, "セキュリティ・運用上の注意", 6)
    bullets(doc, s["notes"])
    callout_box(doc, "共通の注意", [
        "APIトークン/シークレットは秘匿情報です。共有・コミットせず、最小権限で発行してください。",
        "作成・更新・削除などの書き込み操作は、参照のみで検証してから段階的に許可してください。",
    ], fill="FEF2F2", border="DC2626", tcol=RGBColor(0xB9, 0x1C, 0x1C))

    h1(doc, "参考リンク", 7)
    for label, url in s["refs"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        run(p, "・" + label + ": ", 10)
        run(p, url, 9.5, color=ACCENT)

    path = os.path.join(OUT, s["file"] + ".docx")
    doc.save(path)
    return path


def build_readme(docx_paths):
    lines = ["# 営業ツール × Claude 連携マニュアル",
             "",
             "HubSpot の記事『国内の主要SFAツール10選＆特徴や選び方を解説』",
             "(https://blog.hubspot.jp/sales/sfa-comparison) で紹介されている主要SFA/CRMを対象に、",
             "各サービスと Claude / MCP を連携する手順・活用例をまとめたマニュアルです。",
             "**Word(.docx) と PDF をサービス別**に用意しています。",
             "",
             "## 構成",
             "",
             "| 区分 | ファイル(.docx / .pdf) | 内容 |",
             "|---|---|---|",
             "| 共通編 | `00_共通編_Claude×MCP連携の基礎` | MCPの基礎・3つの連携方式・準備・早見表・セキュリティ |"]
    for s in SERVICES:
        lines.append(f"| {s['id']} | `{s['file']}` | {s['name']}({s['vendor']})― {s['tagline']} |")
    lines += [
        "",
        "## 画面図について",
        "",
        "実画面はネットワーク制限・ログイン要件のため取得できないため、各図は操作手順を再現した"
        "**イラスト(モックアップ)**です。各図および本文中に**「実スクリーンショット差込欄」**(オレンジ破線枠)を"
        "設けています。実環境で同じ操作を行い、該当箇所に実画面を貼り付けてご利用ください。",
        "",
        "## 図の見方",
        "",
        "- **赤枠** … 操作対象(クリック/入力する箇所)",
        "- **番号バッジ ①②③** … 操作順序(右側の注釈に対応)",
        "- **オレンジ破線枠** … 実スクリーンショットの差込位置",
        "",
        "## 注意",
        "",
        "各サービスのMCP対応状況・API仕様・画面構成・料金プランは変更される場合があります。",
        "導入前に必ず各社の公式ドキュメントで最新情報をご確認ください。",
        "",
        "## ファイル生成方法(再生成する場合)",
        "",
        "```bash",
        "pip install python-docx Pillow",
        "python tools/build.py   # 図版生成→docx組版→PDF変換まで一括",
        "```",
        "",
        "## 生成物一覧",
        "",
    ]
    for p in docx_paths:
        base = os.path.basename(p)
        name = os.path.splitext(base)[0]
        lines.append(f"- `manuals/{name}.docx` / `manuals/{name}.pdf`")
    with open(os.path.join(OUT, "README.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


def main():
    docx_paths = []
    print("== 共通編 ==")
    docx_paths.append(build_common())
    for s in SERVICES:
        print(f"== {s['id']} {s['name']} ==")
        docx_paths.append(build_service(s))
    build_readme(docx_paths)
    print("Word(.docx) 生成完了: ", len(docx_paths), "件")
    print("PDFは pdfbuild.py で生成します。")


if __name__ == "__main__":
    main()
